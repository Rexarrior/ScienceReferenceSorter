import docx
from os import path
import os
import re

CURDIR = path.dirname(path.abspath(__file__))
DOC_NAME = path.join(CURDIR, "diplom.docx")
OLDREF_NAME = path.join(CURDIR, "oldref.txt")
MATCHREF_NAME = path.join(CURDIR, "newref.txt")
REFLIST_TEXT = "Список литературы"

regex_ref_in_text = re.compile(r"\[[\d;, ]+\]")
regex_is_cyrylic = re.compile(r"[а-яА-Я]")


def get_text_from_doc(doc):
    text = []
    for p in doc.paragraphs:
        text.append(p.text)
    return "\n".join(text)


def get_reflist_from_text(text):
    reflist_start = text.find( REFLIST_TEXT)
    reflist_text = text[reflist_start::]
    lines = reflist_text.split("\n")
    reflist = []
    for i in range(1, len(lines)):
        reflist.append([i, lines[i]])
    return reflist


def reflist_sort_func(refline):
    forsort = refline[1].strip()
    if (regex_is_cyrylic.search(forsort[:10])):
        return "1" + forsort
    return forsort


def sort_reflist_alphabetically(reflist):
    sorted_reflist = reflist.copy()
    sorted_reflist = sorted(sorted_reflist, key=reflist_sort_func)
    refdic = {}
    for i in range(len(sorted_reflist)):
        refline = sorted_reflist[i]
        oldnum = refline[0]
        refline[0] = i + 1 #newnum
        refline_new = [oldnum]
        refline_new.extend(refline)    
        refdic[oldnum] = refline_new
    return refdic


def replace_ref_in_brackets(text, bracket_match, reflist_sorted):
    brackets = bracket_match.group()[1 : len(bracket_match.group()) - 1]
    brackets = brackets.replace(",", ";")
    numbers = brackets.split(";")
    numbers = [num.strip() for num in numbers]
    newbrackets = "[" +  "; ".join([str(reflist_sorted[int(num)][1]) for num in numbers]) +"]"
    newtext = ""
    if (bracket_match.start() > 0):
        newtext += text[:bracket_match.start()]
    newtext += newbrackets
    if (bracket_match.end() < len(text)):
        newtext += text[bracket_match.end():]
    print(f"Заменил [{brackets}] на {newbrackets}\n")
    return newtext


def replace_ref_in_paragraph (paragraph, reflist_sorted):
    text = paragraph.text
    match = regex_ref_in_text.search(text)
    last_end = 0
    while not(match is None):
        if match.end() <= last_end:
            continue
        text = replace_ref_in_brackets(text, match, reflist_sorted)
        last_end = match.end()
        match = regex_ref_in_text.search(text, match.end())
    return text

def replace_ref_in_doc(doc, reflist_sorted):
    for p in doc.paragraphs[:len(doc.paragraphs)-1]: # ignore reference list in the tail of document
        newtext = replace_ref_in_paragraph(p, reflist_sorted)
        p.text = newtext


def find_ref_in_brackets(text, bracket_match):
    brackets = bracket_match.group()[1 : len(bracket_match.group()) - 1]
    brackets = brackets.replace(",", ";")
    numbers = brackets.split(";")
    numbers = [int(num.strip()) for num in numbers]
    return numbers


def find_ref_in_paragraph (paragraph):
    refs = []
    text = paragraph.text
    match = regex_ref_in_text.search(text)
    last_end = 0
    while not(match is None):
        if match.end() <= last_end:
            continue
        refs.extend(find_ref_in_brackets(text, match))
        last_end = match.end()
        match = regex_ref_in_text.search(text, match.end())
    return refs

def find_ref_in_doc(doc):
    used_set = set()
    for p in doc.paragraphs[:len(doc.paragraphs)-1]: # ignore reference list in the tail of document
        used_set = used_set.union(find_ref_in_paragraph(p))
    return used_set
  
def save_in_file(filename, text):
    if path.exists(filename):
        os.remove(filename)
    
    with open(filename, "wt", encoding="utf8") as f:
        f.write(text)

def reflist_to_text_debug(reflist):
    text = "\n".join([")".join([str(listelem) for listelem in reflist[i]]) for i in range(len(reflist))])
    return text




def remove_unused_refs(reflist, used_set):
    return [ref for ref in reflist if ref[0] in used_set]

def second_sort_func(elem):
    return elem[1]


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def replace_reflist_paragraph(doc, sorted_reflist):
    ref_list = [reflist_sorted[key] for key in reflist_sorted.keys()]
    ref_list = sorted(ref_list, key=second_sort_func)
    
    index = 0
    for i in range(len(doc.paragraphs)):
        if (REFLIST_TEXT in doc.paragraphs[i].text):
           index = i
           continue
    style = doc.paragraphs[index].runs[0].style
    for i in range( len(doc.paragraphs)-1, index, -1):
        delete_paragraph(doc.paragraphs[i])
    for refline in ref_list:
        doc.add_paragraph()
        doc.paragraphs[-1].add_run(refline[2])
        doc.paragraphs[-1].runs[0].style = style



if __name__ == "__main__":
    print("Привет! Я работаю...\n")
    filename = DOC_NAME
    doc = docx.Document(filename)
    print(f"Открыл {filename}\n")
    text = get_text_from_doc(doc)
    reflist = get_reflist_from_text(text)
    print(f"Прочитал список. В нем {len(reflist)} позиций.\n")
    oldreftext = reflist_to_text_debug(reflist)
    save_in_file(OLDREF_NAME ,oldreftext)
    print(f"Сохранил старый список в файл {OLDREF_NAME}\n")
    used_refs_set = find_ref_in_doc(doc)
    print(f"Нашел все используемые ссылки. Их ровно {len(used_refs_set)}")
    reflist = remove_unused_refs(reflist, used_refs_set)
    reflist_sorted = sort_reflist_alphabetically(reflist)
    print("Отсортировал по алфавиту\n")
    newref_text = reflist_to_text_debug([reflist_sorted[key] for key in reflist_sorted.keys()])
    save_in_file(MATCHREF_NAME, newref_text)
    print(f"Сохранил новый список в файл {MATCHREF_NAME}\n")
    replace_reflist_paragraph(doc, reflist_sorted)
    print("Заменил список литературы на новый.\nВыполняю замену ссылок\n")
    replace_ref_in_doc(doc, reflist_sorted)
    print("Закончил заменять ссылки\n")
    newfilename = path.join(path.dirname(filename), "processed_" + path.split(filename)[-1])
    doc.save(newfilename)
    print(f"Сохранил в {newfilename}\n")
    print("Пока! Было весело!")

